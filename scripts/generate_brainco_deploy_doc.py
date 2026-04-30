"""Generate a Word (.docx) deployment guide for the brainco hand integration.

Usage:
    python scripts/generate_brainco_deploy_doc.py [output_path]

Default output: decoupled_wbc/control/teleop/device/pico/brainco/Brainco_Deploy_Guide.docx
"""

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = (
    REPO_ROOT
    / "decoupled_wbc"
    / "control"
    / "teleop"
    / "device"
    / "pico"
    / "brainco"
    / "Brainco_Deploy_Guide.docx"
)


def _set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_run_font(run, name="Consolas", size=10)


def _add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    run = p.add_run(text)
    _set_run_font(run)


def _add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, bold=bold, italic=italic)


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def _add_table(doc, header, rows, col_widths_inches=None):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        _set_run_font(run, bold=True, size=10)
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            run = p.add_run(str(v))
            _set_run_font(run, size=10)
    if col_widths_inches:
        for col_idx, width in enumerate(col_widths_inches):
            for row in table.rows:
                row.cells[col_idx].width = Inches(width)


def build_doc():
    doc = Document()

    # Margins
    for section in doc.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # ---------- Title page ----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Brainco RevoLimb Hand Integration")
    _set_run_font(run, size=24, bold=True, color=RGBColor(0x1F, 0x3A, 0x5F))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Gear-Sonic — Change Log & Deployment Guide")
    _set_run_font(run, size=14, color=RGBColor(0x40, 0x40, 0x40))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "PICO XRoboToolkit hand-tracking → dex_retargeting → brainco DDS\n"
        "Status: software-validated; pending hardware bring-up (Stages 2–4)"
    )
    _set_run_font(run, size=11, italic=True)

    doc.add_paragraph()

    # ---------- 1. Executive summary ----------
    _add_heading(doc, "1. Executive summary", level=1)
    _add_para(
        doc,
        "This release adds Brainco RevoLimb dexterous-hand control to the "
        "Gear-Sonic teleop stack. The PICO headset's OpenXR hand-tracking "
        "stream now drives both:",
    )
    _add_bullet(
        doc,
        "the existing arm IK pipeline, replacing button-encoded synthetic "
        "fingertip transforms with real anatomical fingertip positions; and",
    )
    _add_bullet(
        doc,
        "a new parallel pipeline that retargets the hand pose to brainco "
        "joint angles and publishes them on the rt/brainco/{left,right}/cmd "
        "DDS topics.",
    )
    _add_para(
        doc,
        "All software-only validation passes (compile, 18 unit/e2e tests, "
        "smoke test). Hardware-side validation requires a connected PICO and "
        "powered brainco hands and is described in §6 of this document.",
    )

    # ---------- 2. What changed ----------
    _add_heading(doc, "2. What changed", level=1)

    _add_heading(doc, "2.1 New module: decoupled_wbc...pico.brainco", level=2)
    _add_table(
        doc,
        ["File", "Purpose"],
        [
            ("__init__.py", "Lazy-loads BraincoController so dev environments without unitree_sdk2py can still import the bridge math."),
            ("brainco_bridge.py", "Pure-NumPy frame transform: xrobotoolkit (26, 7) hand state → (25, 3) keypoint cloud in unitree-hand frame; shared-array helpers."),
            ("hand_retargeting.py", "BraincoHandRetargeting wrapper around dex_retargeting. Normalises the union-style brainco.yml (DexPilot/vector key suffixes) into the canonical RetargetingConfig shape."),
            ("robot_hand_brainco.py", "BraincoController: subprocess control loop that retargets, normalises to [0, 1] per the brainco API, and publishes on DDS. Includes init timeout, warmup gate, and tracking-loss revert."),
            ("assets/brainco_hand/", "URDFs (left + right), STL meshes, and brainco.yml — all copied verbatim from xr_teleoperate."),
            ("README.md", "In-tree integration doc with derivation, runbook, and failure-mode table."),
        ],
        col_widths_inches=[2.2, 4.6],
    )

    _add_heading(doc, "2.2 New runners and tests", level=2)
    _add_table(
        doc,
        ["File", "Purpose"],
        [
            ("control/teleop/main/run_brainco_teleop.py", "Standalone CLI (no arm IK), exposes --state-timeout-s and --warmup-frames flags."),
            ("control/teleop/main/smoke_test_brainco.py", "Hardware-free pre-deploy check: imports + dex_retargeting URDF wiring + bridge math on synthetic input."),
            ("tests/control/teleop/brainco/test_brainco_bridge.py", "14 unit tests covering rotation correctness, palm-drop, frame anchor, and shared-array roundtrip."),
            ("tests/control/teleop/brainco/test_brainco_retarget_e2e.py", "4 e2e tests, gated on dex_retargeting being installed: YAML loads, retarget runs, output sensitivity, motor-angle ranges."),
        ],
        col_widths_inches=[3.0, 3.8],
    )

    _add_heading(doc, "2.3 Modified files", level=2)
    _add_table(
        doc,
        ["File", "Change"],
        [
            ("decoupled_wbc/control/teleop/streamers/pico_streamer.py",
             "Constructor gained enable_brainco, brainco_fps, brainco_state_timeout_s, brainco_warmup_frames. _generate_finger_data now fills (25, 4, 4) with real wrist-anchored hand keypoints (in meters). Falls back to button-encoded synthetic signal only when hand-tracking is_active=0."),
            ("decoupled_wbc/control/teleop/device/pico/xr_client.py",
             "Hand-tracking docstring corrected from 27×7 to 26×7, with the Khronos OpenXR XR_EXT_hand_tracking joint enum and pose convention spelled out."),
            ("gear_sonic/pyproject.toml",
             "Added dex-retargeting and pyyaml to the [teleop] optional-dependency group, so install_pico.sh picks them up automatically."),
        ],
        col_widths_inches=[3.2, 3.6],
    )

    # ---------- 3. Why these changes ----------
    _add_heading(doc, "3. Why these changes", level=1)
    _add_para(
        doc,
        "The primary motivation is replacing fake fingertip data with real "
        "hand tracking. The previous _generate_finger_data path encoded "
        "trigger/grip button combos as synthetic (1, 0, 0) fingertip "
        "transforms — useful as a pinch fallback but blind to the user's "
        "actual fingers. With hand tracking already available from "
        "xrobotoolkit, the synthetic path was wasting a real signal.",
    )
    _add_para(
        doc,
        "The second motivation is enabling Brainco RevoLimb hand control. "
        "Reusing the same xrobotoolkit hand state for both pipelines avoids "
        "a second SDK connection and keeps the data flow auditable.",
    )

    _add_heading(doc, "3.1 Data flow", level=2)
    _add_code(
        doc,
        "PICO headset (OpenXR app)\n"
        "  └─ XRoboToolkit-PC-Service\n"
        "       └─ xrobotoolkit_sdk\n"
        "            └─ XrClient.get_hand_tracking_state()        [main process]\n"
        "                 ├─► hand_state_to_unitree_keypoints()  → (25, 3) keypoints\n"
        "                 │      │\n"
        "                 │      ├─► push_keypoints_to_shared()  → Array('d', 75)\n"
        "                 │      │      │\n"
        "                 │      │      └─► BraincoController._control_loop  [subprocess]\n"
        "                 │      │             dex_retargeting → normalise → DDS\n"
        "                 │      │             rt/brainco/{left,right}/cmd\n"
        "                 │      │\n"
        "                 │      └─► fingertips[:, :3, 3]   → (25, 4, 4)\n"
        "                 │             │\n"
        "                 │             └─► G1GripperInverseKinematicsSolver\n"
        "                 │                    thumb-to-tip distances → arm pose\n"
        "                 └─► Brainco RevoLimb hands (via DDS)",
    )

    _add_heading(doc, "3.2 Coordinate-frame correctness", level=2)
    _add_para(
        doc,
        "dex_retargeting expects wrist-relative finger vectors in the brainco "
        "URDF base_link frame. We arrive there from xrobotoolkit's OpenXR "
        "hand-tracking output via a single composed rotation:",
    )
    _add_code(
        doc,
        "R_HAND_TO_UNITREE = T_TO_UNITREE_HAND @ R_ROBOT_OPENXR\n"
        "                  = [[0, 1, 0],\n"
        "                     [0, 0, 1],\n"
        "                     [1, 0, 0]]\n"
        "\n"
        "delta_unitree = R_HAND_TO_UNITREE @ wrist_R_openxr.T @ (tip - wrist)_world",
    )
    _add_para(
        doc,
        "Both T_TO_UNITREE_HAND and R_ROBOT_OPENXR are copied verbatim from "
        "xr_teleoperate's televuer pipeline. The composed matrix is "
        "orthogonal with det = 1 (verified in test_R_HAND_TO_UNITREE_*). "
        "The full algebraic derivation is in the in-tree README §6.",
    )

    # ---------- 4. Safety features ----------
    _add_heading(doc, "4. Safety features", level=1)
    _add_para(
        doc,
        "BraincoController has three runtime guards. Each is exposed as a "
        "kwarg on BraincoController and PicoStreamer, and as a CLI flag on "
        "run_brainco_teleop.py.",
    )

    _add_table(
        doc,
        ["Guard", "Default", "Behaviour"],
        [
            ("DDS state-init timeout (state_timeout_s)", "5.0 s",
             "If rt/brainco/{left,right}/state isn't publishing, raises TimeoutError instead of hanging the streamer forever. Set to 0 to wait forever (legacy)."),
            ("Warmup gate (warmup_frames)", "10 frames",
             "First N consecutive valid input frames are NOT retargeted; controller publishes q=0 (open hand). Blocks cold-start xrobotoolkit junk from slamming the hand closed."),
            ("Tracking-loss revert", "always on",
             "If input becomes all-zero or non-finite after warmup, controller drops back to open hand and re-arms the warmup counter. One noisy frame costs you warmup_frames/fps seconds of \"open hand\" but eliminates the slam-shut failure mode."),
        ],
        col_widths_inches=[2.0, 1.0, 3.8],
    )

    # ---------- 5. Validation status ----------
    _add_heading(doc, "5. Validation status (what's done)", level=1)
    _add_table(
        doc,
        ["Stage", "Validation", "Status"],
        [
            ("0", "Every brainco source file compiles cleanly (ast.parse + compile, no warnings).", "PASS"),
            ("1a", "dex_retargeting can build BraincoHandRetargeting from the in-tree URDFs and brainco.yml. Verified with python-docx 3.10 venv mirroring install_pico.sh.", "PASS"),
            ("1b", "14 unit tests on bridge math: rotation correctness, palm-drop, frame anchor, axis mapping, distance preservation, shared-array roundtrip.", "PASS (14/14)"),
            ("1c", "4 e2e tests with real dex_retargeting: YAML normalisation, retarget runs, response to input changes, motor-angle ranges within brainco spec.", "PASS (4/4)"),
            ("1d", "smoke_test_brainco.py: imports + asset wiring + bridge math, EXIT 0 on dev box.", "PASS"),
            ("2", "Verify xrobotoolkit emits OpenXR-spec hand-joint data on real PICO firmware (5-line print script).", "Pending hardware"),
            ("3", "Brainco hands powered, parked safe — verify warmup, mirror open/close, pinch.", "Pending hardware"),
            ("4", "Full integrated PicoStreamer(enable_brainco=True) with arm IK.", "Pending hardware"),
        ],
        col_widths_inches=[0.6, 4.7, 1.5],
    )

    _add_heading(doc, "5.1 Notable bug caught and fixed during validation", level=2)
    _add_para(
        doc,
        "brainco.yml uses xr_teleoperate's union-style key naming "
        "(target_link_human_indices_dexpilot vs ..._vector). Current "
        "dex_retargeting (≥0.4) rejects suffixed keys with a TypeError. "
        "Without the fix, BraincoController() would have crashed on first "
        "construction. Resolved in hand_retargeting._normalize_section() and "
        "regression-guarded by test_retargeter_loads_with_brainco_yaml.",
    )

    # ---------- 6. Deployment runbook ----------
    _add_heading(doc, "6. Deployment runbook", level=1)
    _add_para(
        doc,
        "Run these stages in order. Do not skip ahead until each passes. "
        "Every command assumes you are at the repo root (Gear-sonic/GR00T-WholeBodyControl).",
    )

    _add_heading(doc, "6.1 Install", level=2)
    _add_code(
        doc,
        "bash install_scripts/install_pico.sh\n"
        "source .venv_teleop/bin/activate",
    )
    _add_para(
        doc,
        "This sets up a Python 3.10 venv, installs gear_sonic[teleop] "
        "(which now includes dex-retargeting and pyyaml), and builds "
        "unitree_sdk2py and xrobotoolkit_sdk from "
        "external_dependencies/.",
    )

    _add_heading(doc, "6.2 Stage 0 — hardware-free smoke test", level=2)
    _add_code(doc, "python -m decoupled_wbc.control.teleop.main.smoke_test_brainco")
    _add_para(
        doc,
        "All three stages should print [OK]. Exit 0 = ready to proceed. On "
        "the deployment box every line must be [OK]; if you see [DEP] lines, "
        "the runtime deps are not installed and you skipped step 6.1.",
    )

    _add_heading(doc, "6.3 Stage 1 — unit tests", level=2)
    _add_code(doc, "pytest decoupled_wbc/tests/control/teleop/brainco/")
    _add_para(doc, "Should report 18 passed.")

    _add_heading(doc, "6.4 Stage 2 — PICO connected, brainco hands UNPOWERED", level=2)
    _add_para(
        doc,
        "Goal: confirm xrobotoolkit emits OpenXR-spec hand-joint data on "
        "this firmware. Connect the PICO, launch the XRoboToolkit Unity app, "
        "enable hand tracking, then run:",
    )
    _add_code(
        doc,
        "python - <<'EOF'\n"
        "import xrobotoolkit_sdk as xrt; xrt.init()\n"
        "import time\n"
        "while True:\n"
        "    s = xrt.get_left_hand_tracking_state()\n"
        "    print('palm_z=%.3f wrist_z=%.3f tip_index_z=%.3f' %\n"
        "          (s[0][2], s[1][2], s[10][2]))\n"
        "    time.sleep(0.5)\n"
        "EOF",
    )
    _add_para(
        doc,
        "Hold your left hand at chest height with fingers up. Verify "
        "wrist_z < tip_index_z (fingertip is higher than wrist). If "
        "wrist_z > tip_index_z, the firmware uses a non-standard joint "
        "ordering and the spec assumption needs review before continuing.",
    )

    _add_heading(doc, "6.5 Stage 3 — brainco hands powered, parked safe", level=2)
    _add_para(
        doc,
        "Power the hands in a position where a sudden close cannot damage "
        "them or the robot.",
    )
    _add_code(
        doc,
        "python -m decoupled_wbc.control.teleop.main.run_brainco_teleop \\\n"
        "    --network-interface eth0",
    )
    _add_para(doc, "Expected log sequence:", bold=True)
    _add_bullet(doc, "xrobotoolkit SDK initialised.")
    _add_bullet(doc, "Initialising BraincoController")
    _add_bullet(doc, "[BraincoController] DDS state ready  (within 5 s, or TimeoutError fires)")
    _add_bullet(doc, "BraincoController control process started")
    _add_bullet(doc, "[BraincoController] warmup complete after 10 valid frames; publishing retargeted commands  (within ~100 ms of seeing your hand)")

    _add_para(doc, "Then test interactively:", bold=True)
    _add_bullet(doc, "Hold hand still and open — brainco hand holds open.")
    _add_bullet(doc, "Slowly close to a fist — brainco hand mirrors.")
    _add_bullet(doc, "Pinch thumb-to-index — motor 0 (thumb) and motor 2 (index) close.")
    _add_bullet(doc, "Move hand out of camera FOV — log shows tracking lost; reverting to open hand.")
    _add_bullet(doc, "Move back in — warmup completes, retargeting resumes.")

    _add_heading(doc, "6.6 Stage 4 — full integrated teleop", level=2)
    _add_para(doc, "Only after Stages 0–3 pass:")
    _add_code(
        doc,
        "from decoupled_wbc.control.teleop.streamers.pico_streamer import PicoStreamer\n"
        "streamer = PicoStreamer(enable_brainco=True)\n"
        "# rest of your existing teleop main loop unchanged",
    )
    _add_para(
        doc,
        "Optional kwargs: brainco_fps (default 100), brainco_state_timeout_s "
        "(default 5.0), brainco_warmup_frames (default 10).",
    )

    # ---------- 7. Failure modes ----------
    _add_heading(doc, "7. Failure modes & how to debug", level=1)
    _add_table(
        doc,
        ["Symptom", "Likely cause", "Fix"],
        [
            ("TimeoutError: no DDS state on rt/brainco/...",
             "Brainco hands not powered, or wrong DDS interface.",
             "Confirm hands powered, on the same network. Check --network-interface. Or pass state_timeout_s=0 to wait forever."),
            ("Hand opens to safe pose then never closes",
             "Warmup gate keeps resetting — transient is_active=0.",
             "Move hand into camera FOV. If still failing, lower --warmup-frames (default 10) once you trust the source."),
            ("Brainco fingers move but inverted/rotated",
             "Frame composition wrong.",
             "Check R_HAND_TO_UNITREE in brainco_bridge.py against §3.2. Most likely fix is a sign flip on one column."),
            ("Single finger lags / never closes",
             "Brainco firmware motor order differs from BraincoLeftJointIndex.",
             "Compare against brainco-hz.com/docs/revolimb-hand/product/parameters.html. Adjust IntEnum in robot_hand_brainco.py."),
            ("_generate_finger_data falls back to button mode",
             "xr_client.get_hand_tracking_state returned None (is_active=0).",
             "Hand outside FOV or occlusion. Move into camera view."),
            ("Distance threshold never triggers in arm IK",
             "Keypoints in non-meter units.",
             "Print np.linalg.norm(positions[4] - positions[9]) at fist; should be ~3 cm."),
            ("ImportError: dex_retargeting",
             "Teleop venv not installed or not active.",
             "bash install_scripts/install_pico.sh && source .venv_teleop/bin/activate."),
            ("TypeError: ...target_link_human_indices_dexpilot...",
             "Old hand_retargeting.py without _normalize_section.",
             "Pull the latest brainco/hand_retargeting.py — this fix is in the deploy package."),
        ],
        col_widths_inches=[2.0, 2.2, 2.6],
    )

    # ---------- 8. Quick reference ----------
    _add_heading(doc, "8. Quick reference", level=1)

    _add_heading(doc, "8.1 DDS topics", level=2)
    _add_table(
        doc,
        ["Direction", "Topic", "Type"],
        [
            ("Subscribe", "rt/brainco/left/state",  "MotorStates_"),
            ("Subscribe", "rt/brainco/right/state", "MotorStates_"),
            ("Publish",   "rt/brainco/left/cmd",    "MotorCmds_"),
            ("Publish",   "rt/brainco/right/cmd",   "MotorCmds_"),
        ],
        col_widths_inches=[1.2, 2.4, 1.6],
    )

    _add_heading(doc, "8.2 Brainco motor mapping", level=2)
    _add_table(
        doc,
        ["Motor", "Joint", "Range (rad)"],
        [
            ("0", "thumb metacarpal", "[0, 1.52]"),
            ("1", "thumb proximal (aux)", "[0, 1.05]"),
            ("2", "index proximal", "[0, 1.47]"),
            ("3", "middle proximal", "[0, 1.47]"),
            ("4", "ring proximal", "[0, 1.47]"),
            ("5", "pinky proximal", "[0, 1.47]"),
        ],
        col_widths_inches=[0.8, 2.6, 1.6],
    )
    _add_para(
        doc,
        "Published cmd values are normalised to [0, 1] by _normalize() — "
        "0 = open, 1 = fully closed. Reference: "
        "brainco-hz.com/docs/revolimb-hand/product/parameters.html.",
        italic=True,
    )

    _add_heading(doc, "8.3 OpenXR hand-joint enum (XR_EXT_hand_tracking)", level=2)
    _add_table(
        doc,
        ["Index", "Joint", "Index", "Joint"],
        [
            ("0",  "PALM (dropped)",       "13", "middle proximal"),
            ("1",  "WRIST → kept index 0", "14", "middle intermediate"),
            ("2",  "thumb metacarpal",     "15", "middle distal"),
            ("3",  "thumb proximal",       "16", "ring metacarpal"),
            ("4",  "thumb distal",         "17", "ring proximal"),
            ("5",  "thumb tip → kept 4",   "18", "ring intermediate"),
            ("6",  "index metacarpal",     "19", "ring distal → kept 19"),
            ("7",  "index proximal",       "20", "ring tip"),
            ("8",  "index intermediate",   "21", "little metacarpal"),
            ("9",  "index distal",         "22", "little proximal"),
            ("10", "index tip → kept 9",   "23", "little intermediate"),
            ("11", "middle metacarpal",    "24", "little distal → kept 24"),
            ("12", "middle (skip)",        "25", "little tip"),
        ],
        col_widths_inches=[0.7, 2.6, 0.7, 2.6],
    )
    _add_para(
        doc,
        "After dropping PALM, the bridge produces a (25, 3) cloud whose "
        "wrist sits at index 0 and fingertips at 4 / 9 / 14 / 19 / 24 — "
        "matching the indices referenced by brainco.yml.",
        italic=True,
    )

    # ---------- 9. Sign-off ----------
    _add_heading(doc, "9. Pre-deployment sign-off checklist", level=1)
    _add_bullet(doc, "Smoke test passes on the target machine (Stage 0).")
    _add_bullet(doc, "All 18 unit + e2e tests pass on the target machine (Stage 1).")
    _add_bullet(doc, "PICO Stage-2 print script confirms wrist_z < tip_index_z.")
    _add_bullet(doc, "Stage 3 standalone runner mirrors the operator's hand with no 90°-off rotation.")
    _add_bullet(doc, "Tracking-loss revert log line observed when hand leaves FOV.")
    _add_bullet(doc, "Warmup-complete log line observed within ~150 ms of hand re-entering FOV.")
    _add_bullet(doc, "PicoStreamer(enable_brainco=True) integrated test passes; arm IK pipeline unchanged.")
    _add_bullet(doc, "Operator briefed: hand commands route over rt/brainco/{left,right}/cmd; existing arm IK observes real fingertip distances now (5 cm pinch threshold).")

    _add_para(doc, "")
    _add_para(
        doc,
        "Document generated by scripts/generate_brainco_deploy_doc.py. "
        "If anything in this guide drifts from the source, regenerate.",
        italic=True,
    )

    return doc


def main():
    output = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_OUTPUT
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(str(output))
    print(f"Wrote {output}")


if __name__ == "__main__":
    main()
